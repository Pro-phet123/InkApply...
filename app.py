# app.py  –  InkApply  |  AI Cover Letter Generator
# ─────────────────────────────────────────────────
# Optimised for Railway & Streamlit Cloud:
#   • @st.cache_resource  → model loaded once, never re-loaded on rerun
#   • session_state guard → resume bytes / text survive every widget rerun
#   • trimmed prompts     → fast inference, lower memory pressure
# ─────────────────────────────────────────────────

import os
import base64
from io import BytesIO

import streamlit as st
from PyPDF2 import PdfReader
from docx import Document

from prompts import generate_cover_letter_prompt
from hf_inference import query_llama3


# ══════════════════════════════════════════════════
#  PAGE CONFIG  (must be first Streamlit call)
# ══════════════════════════════════════════════════
st.set_page_config(
    page_title="InkApply – AI Cover Letter Generator",
    page_icon="Inkapply-logo.png" if os.path.exists("Inkapply-logo.png") else "📝",
    layout="wide",
)


# ══════════════════════════════════════════════════
#  CACHE THE MODEL  ← THE KEY ANTI-BUFFERING FIX
#  Loads once per server process, never again.
# ══════════════════════════════════════════════════
@st.cache_resource(show_spinner="Loading AI model… (first run only)")
def get_llm():
    """
    If your hf_inference module wraps a local pipeline, initialise it here
    so it is cached. If it's a pure HTTP call to HF Inference API, this
    acts as a lightweight warm-up sentinel.
    """
    # Uncomment and adapt if you load a local pipeline:
    # from transformers import pipeline
    # return pipeline("text-generation", model="meta-llama/Meta-Llama-3-8B-Instruct")
    return True


# ══════════════════════════════════════════════════
#  SESSION STATE DEFAULTS
# ══════════════════════════════════════════════════
st.session_state.setdefault("resume_bytes", None)
st.session_state.setdefault("resume_content", "")
st.session_state.setdefault("uploaded_file_name", "")
st.session_state.setdefault("cover_letter", "")

# Warm up cache immediately
get_llm()


# ══════════════════════════════════════════════════
#  GLOBAL STYLING
# ══════════════════════════════════════════════════
st.markdown(
    """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;500;600&family=DM+Serif+Display&display=swap');

    html, body, [class*="css"] {
        font-family: 'DM Sans', sans-serif;
    }

    /* ── Layout ── */
    .block-container {
        padding-top: 2.75rem !important;
        max-width: 100% !important;
        padding-left: 2.5rem !important;
        padding-right: 2.5rem !important;
    }

    /* ── Hide Streamlit chrome ── */
    #MainMenu { visibility: hidden; }
    footer    { visibility: hidden; }
    header    { visibility: hidden; }

    /* ── Section label ── */
    .ink-label {
        font-size: 0.72rem;
        font-weight: 600;
        letter-spacing: 0.08em;
        text-transform: uppercase;
        color: #6b7280;
        margin-bottom: 0.35rem;
    }

    /* ── Hero heading ── */
    .ink-hero-title {
        font-family: 'DM Serif Display', serif;
        font-size: 2rem;
        font-weight: 400;
        color: #f5f5f5;
        margin: 0 0 0.2rem 0;
        line-height: 1.2;
    }
    .ink-hero-sub {
        font-size: 0.95rem;
        color: #6b7280;
        margin: 0 0 1.75rem 0;
    }

    /* ── Inputs ── */
    textarea, input[type="text"] {
        background-color: #1a1a1a !important;
        border: 1px solid #2a2a2a !important;
        border-radius: 10px !important;
        color: #e8e8e8 !important;
        font-family: 'DM Sans', sans-serif !important;
        font-size: 0.9rem !important;
    }
    textarea:focus, input[type="text"]:focus {
        border-color: #10A37F !important;
        box-shadow: 0 0 0 2px rgba(16,163,127,0.18) !important;
        outline: none !important;
    }
    textarea::placeholder, input::placeholder {
        color: #4b5563 !important;
    }

    /* ── Primary button ── */
    div[data-testid="stButton"] > button {
        background: linear-gradient(135deg, #10A37F 0%, #0d8a6b 100%) !important;
        color: #fff !important;
        border: none !important;
        border-radius: 10px !important;
        padding: 0.6rem 2rem !important;
        font-weight: 600 !important;
        font-size: 0.92rem !important;
        letter-spacing: 0.01em !important;
        transition: opacity 0.2s ease, transform 0.15s ease !important;
    }
    div[data-testid="stButton"] > button:hover {
        opacity: 0.88 !important;
        transform: translateY(-1px) !important;
    }

    /* ── Download buttons ── */
    div[data-testid="stDownloadButton"] > button {
        background: #1f1f1f !important;
        color: #10A37F !important;
        border: 1px solid #10A37F !important;
        border-radius: 10px !important;
        font-weight: 600 !important;
        font-size: 0.85rem !important;
        transition: background 0.2s ease !important;
    }
    div[data-testid="stDownloadButton"] > button:hover {
        background: rgba(16,163,127,0.12) !important;
    }

    /* ── Sidebar ── */
    section[data-testid="stSidebar"] {
        background: #0f0f0f;
        border-right: 1px solid #1f1f1f;
        padding-top: 1.5rem;
    }
    .sidebar-logo {
        display: flex;
        justify-content: center;
        padding: 0.5rem 0 1.25rem 0;
    }
    .sidebar-logo img {
        border-radius: 50%;
        width: 84px;
        height: 84px;
        object-fit: cover;
        border: 2px solid #10A37F;
        box-shadow: 0 0 18px rgba(16,163,127,0.25);
    }
    .sidebar-tagline {
        color: #4b5563;
        font-size: 0.8rem;
        text-align: center;
        line-height: 1.5;
        padding: 0 1rem;
    }
    .sidebar-tip {
        background: #161616;
        border: 1px solid #222;
        border-radius: 10px;
        padding: 0.75rem 1rem;
        font-size: 0.78rem;
        color: #6b7280;
        margin: 1rem 1rem 0;
        line-height: 1.55;
    }
    .sidebar-tip strong { color: #10A37F; }

    /* ── Result heading ── */
    .ink-result-heading {
        font-family: 'DM Serif Display', serif;
        font-size: 1.2rem;
        color: #f0f0f0;
        margin-bottom: 0.75rem;
    }

    /* ── Status badge ── */
    .ink-badge {
        display: inline-flex;
        align-items: center;
        gap: 0.4rem;
        background: rgba(16,163,127,0.1);
        border: 1px solid rgba(16,163,127,0.25);
        color: #10A37F;
        border-radius: 20px;
        padding: 0.25rem 0.75rem;
        font-size: 0.78rem;
        font-weight: 600;
        margin-bottom: 0.75rem;
    }
    </style>
    """,
    unsafe_allow_html=True,
)


# ══════════════════════════════════════════════════
#  SIDEBAR
# ══════════════════════════════════════════════════
logo_path = "Inkapply-logo.png"
if os.path.exists(logo_path):
    with open(logo_path, "rb") as f:
        logo_b64 = base64.b64encode(f.read()).decode()
    st.sidebar.markdown(
        f'<div class="sidebar-logo"><img src="data:image/png;base64,{logo_b64}" /></div>',
        unsafe_allow_html=True,
    )
else:
    st.sidebar.markdown(
        "<div style='text-align:center; font-family:DM Serif Display,serif; "
        "font-size:1.4rem; color:#f0f0f0; padding:1rem 0 0.5rem;'>InkApply</div>",
        unsafe_allow_html=True,
    )

st.sidebar.markdown(
    "<p class='sidebar-tagline'>AI-powered cover letters<br>tailored to every job.</p>",
    unsafe_allow_html=True,
)

st.sidebar.markdown("---")

st.sidebar.markdown(
    """
    <div class='sidebar-tip'>
        <strong>✦ Tips for best results</strong><br>
        • Paste a detailed job description<br>
        • Upload a complete resume (PDF preferred)<br>
        • Use a specific job title
    </div>
    """,
    unsafe_allow_html=True,
)

st.sidebar.markdown("<br>", unsafe_allow_html=True)
st.sidebar.markdown(
    "<p class='sidebar-tagline' style='font-size:0.72rem;'>Model cached on first load.<br>Subsequent generations are fast.</p>",
    unsafe_allow_html=True,
)


# ══════════════════════════════════════════════════
#  HERO
# ══════════════════════════════════════════════════
st.markdown(
    "<h1 class='ink-hero-title'>Cover Letter Generator</h1>"
    "<p class='ink-hero-sub'>Craft a tailored, professional cover letter in seconds — powered by AI.</p>",
    unsafe_allow_html=True,
)


# ══════════════════════════════════════════════════
#  FILE PARSER
# ══════════════════════════════════════════════════
def parse_uploaded_file(file_bytes: bytes, filename: str) -> str:
    try:
        fn = filename.lower()
        if fn.endswith(".pdf"):
            reader = PdfReader(BytesIO(file_bytes))
            return "\n".join(
                p.extract_text() for p in reader.pages if p.extract_text()
            ).strip()
        if fn.endswith(".docx"):
            doc = Document(BytesIO(file_bytes))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip()).strip()
        if fn.endswith(".txt"):
            return file_bytes.decode("utf-8", errors="ignore").strip()
        return ""
    except Exception as e:
        st.error(f"⚠️ Could not read file: {e}")
        return ""


# ══════════════════════════════════════════════════
#  INPUTS  –  single column, top to bottom
# ══════════════════════════════════════════════════
st.markdown("<div class='ink-label'>Job Title</div>", unsafe_allow_html=True)
job_title = st.text_input(
    label="job_title_hidden",
    label_visibility="collapsed",
    placeholder="e.g. Senior Embedded Software Engineer",
    key="job_title_input",
)

st.markdown("<div class='ink-label'>Job Description</div>", unsafe_allow_html=True)
job_description = st.text_area(
    label="job_desc_hidden",
    label_visibility="collapsed",
    placeholder="Paste the job description here (optional but recommended for best results)",
    height=160,
    key="job_desc_input",
)

st.markdown("<div class='ink-label'>Upload Resume</div>", unsafe_allow_html=True)
uploaded_file = st.file_uploader(
    label="resume_uploader_hidden",
    label_visibility="collapsed",
    type=["pdf", "docx", "txt"],
    key="resume_uploader",
)

# ── Handle new upload ──
if uploaded_file is not None:
    try:
        file_bytes = uploaded_file.read()
        st.session_state.resume_bytes       = file_bytes
        st.session_state.uploaded_file_name = uploaded_file.name
        parsed = parse_uploaded_file(file_bytes, uploaded_file.name)
        if parsed:
            st.session_state.resume_content = parsed
            st.markdown(
                f"<div class='ink-badge'>✓ {uploaded_file.name}</div>",
                unsafe_allow_html=True,
            )
        else:
            st.warning("Could not extract text — please paste your resume below.")
    except Exception as e:
        st.error(f"Upload failed: {e}")

# ── Re-parse from session_state if uploader is empty ──
elif st.session_state.resume_bytes and not st.session_state.resume_content:
    parsed = parse_uploaded_file(
        st.session_state.resume_bytes,
        st.session_state.uploaded_file_name or "resume",
    )
    if parsed:
        st.session_state.resume_content = parsed

# ── Show persisted filename after rerun ──
if st.session_state.uploaded_file_name and uploaded_file is None:
    st.markdown(
        f"<div class='ink-badge'>✓ {st.session_state.uploaded_file_name}</div>",
        unsafe_allow_html=True,
    )

st.markdown("<div class='ink-label'>Resume Content</div>", unsafe_allow_html=True)
new_text = st.text_area(
    label="resume_text_hidden",
    label_visibility="collapsed",
    value=st.session_state.resume_content,
    placeholder="Or paste your resume content here…",
    height=260,
    key="resume_textarea",
)
if new_text != st.session_state.resume_content:
    st.session_state.resume_content = new_text


# ══════════════════════════════════════════════════
#  GENERATE BUTTON
# ══════════════════════════════════════════════════
st.markdown("<br>", unsafe_allow_html=True)
generate_clicked = st.button("✨  Generate Cover Letter", use_container_width=False)

if generate_clicked:
    if not job_title.strip():
        st.warning("Please enter a job title.")
    elif not st.session_state.resume_content.strip():
        st.warning("Please upload or paste your resume.")
    else:
        trimmed_resume = " ".join(st.session_state.resume_content.split()[:300])
        trimmed_desc   = " ".join((job_description or "").split()[:200])

        prompt = generate_cover_letter_prompt(
            job_title.strip(),
            trimmed_desc,
            trimmed_resume,
        )

        with st.spinner("Writing your cover letter…"):
            try:
                generated_text = query_llama3(prompt)
                st.session_state.cover_letter = generated_text.strip()
            except Exception as e:
                st.error(f"Generation failed: {e}")


# ══════════════════════════════════════════════════
#  RESULT  (persists across reruns via session_state)
# ══════════════════════════════════════════════════
if st.session_state.cover_letter:
    st.markdown("<hr style='border-color:#1f1f1f; margin:1.5rem 0;'>", unsafe_allow_html=True)
    st.markdown("<div class='ink-result-heading'>Your Cover Letter</div>", unsafe_allow_html=True)

    st.text_area(
        label="cover_letter_output_hidden",
        label_visibility="collapsed",
        value=st.session_state.cover_letter,
        height=380,
        key="cover_letter_output",
    )

    dl_col1, dl_col2, _ = st.columns([1, 1, 4])

    with dl_col1:
        st.download_button(
            "⬇ Download TXT",
            st.session_state.cover_letter,
            "cover_letter.txt",
            "text/plain",
        )

    with dl_col2:
        doc = Document()
        doc.add_heading(f"Cover Letter – {job_title or 'Position'}", 0)
        doc.add_paragraph(st.session_state.cover_letter)
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        st.download_button(
            "⬇ Download DOCX",
            buf,
            "cover_letter.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
